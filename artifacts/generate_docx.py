import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Inject placeholders
    placeholders = {
        '7.1 Admin Dashboard Screen': '[[Insert fig01_admin_dashboard.png here (Full screen, capturing the KPI cards and side menu)]]',
        '7.2 Artisan Profile & Resume Editor': '[[Insert fig02_artisan_profile.png here (Main content area, showing the resume form and preview)]]',
        '7.3 Artisan Product Management Screen': '[[Insert fig03_product_management.png here (Main content area, capturing the product list table)]]',
        '7.4 Public Marketplace Listing Page': '[[Insert fig04_marketplace.png here (Full screen, showing filters and product grid)]]',
        '7.5 Product Detail Page': '[[Insert fig05_product_detail.png here (Full screen, showing the main image, details, and artisan card)]]',
        '7.6 Wishlist Page': '[[Insert fig06_wishlist.png here (Main content area, showing the list of saved items)]]',
        '8.1 شاشة لوحة تحكم الإدارة': '[[إدراج fig01_admin_dashboard.png هنا]]',
        '8.2 محرر ملف الحرفي والسيرة الذاتية': '[[إدراج fig02_artisan_profile.png هنا]]',
        '8.3 شاشة إدارة منتجات الحرفي': '[[إدراج fig03_product_management.png هنا]]',
        '8.4 صفحة قائمة السوق العام': '[[إدراج fig04_marketplace.png هنا]]',
        '8.5 صفحة تفاصيل المنتج': '[[إدراج fig05_product_detail.png هنا]]',
        '8.6 صفحة قائمة الرغبات': '[[إدراج fig06_wishlist.png هنا]]'
    }

    for key, placeholder in placeholders.items():
        # Find the line with the key, then insert the placeholder before the caption
        pattern = re.compile(rf"(### {re.escape(key)}.*?- \*\*Suggested Caption:\*\*.*?)(?=\n)", re.DOTALL)
        content = re.sub(
            rf"(### {re.escape(key)}.*?)(?=- \*\*Suggested Caption:\*\*|- \*\*التعليق المقترح:\*\*)",
            rf"\1\n\n{placeholder}\n\n",
            content,
            flags=re.DOTALL
        )

    doc = Document()
    
    # Add Arabic compatible style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
        
        is_arabic = any('\u0600' <= c <= '\u06FF' for c in line)
        
        if line.startswith('# '):
            p = doc.add_heading(line[2:].replace('*', ''), 0)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:].replace('*', ''), 1)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:].replace('*', ''), 2)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:].replace('**', ''), style='List Bullet')
        else:
            p = doc.add_paragraph(line.replace('**', '').replace('*', ''))
            
        # Right to left alignment for Arabic
        if is_arabic and p:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(docx_path)

if __name__ == '__main__':
    create_docx('e:/RedOasisArtisan/artifacts/thesis_documentation.md', 'e:/RedOasisArtisan/artifacts/RedOasisArtisan_thesis_documentation.docx')
    print('DOCX generated successfully.')
